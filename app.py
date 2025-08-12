"""Streamlit OCR app with fine-tuning options and reports."""
from __future__ import annotations

import io
import logging
import os
import tempfile
import time
from typing import Any, Dict, List

import numpy as np
import streamlit as st
from PIL import Image, ImageDraw

import easyocr

import layout_heuristics as heur


# ---------------------------------------------------------------------------
# Configuration helpers
# ---------------------------------------------------------------------------

def build_default_config() -> Dict[str, Any]:
    """Return default configuration values for the OCR pipeline."""
    return {
        "dpi": 300,
        "languages": ["de", "en"],
        "gpu": False,
        "column_detection": True,
        "max_columns": 2,
        "heading_threshold_h2": 1.4,
        "heading_threshold_h1": 1.8,
        "heading_extra_rules": {"centered": True, "all_caps": True, "big_gap": True},
        "paragraph_merge_gap": 1.2,
        "indent_tolerance": 0.04,
        "hyphen_merge": True,
        "output_format": "docx",
        "text_alignment": "justify",
        "keep_line_breaks": False,
        "debug_overlay": False,
        "custom_regex_cleanup": [],
    }


def render_controls(config: Dict[str, Any]) -> Dict[str, Any]:
    """Render Streamlit controls and return updated configuration."""
    st.subheader("Grundeinstellungen")
    dpi = st.number_input(
        "DPI",
        min_value=72,
        max_value=600,
        value=config["dpi"],
        help="Auflösung für die PDF-Bild-Konvertierung.",
    )
    if dpi > 450:
        st.warning("Sehr hohe DPI können die Laufzeit stark erhöhen.")
    languages = st.multiselect(
        "Sprachen",
        ["de", "en", "fr", "es", "it"],
        default=config["languages"],
        help="OCR-Sprachen",
    )
    gpu = st.checkbox("GPU verwenden", value=config["gpu"], help="Aktiviere GPU falls vorhanden")
    try:
        import torch

        has_gpu = torch.cuda.is_available()
    except Exception:
        has_gpu = False
    if gpu and not has_gpu:
        st.warning("Keine GPU verfügbar")
    output_format = st.selectbox(
        "Ausgabeformat",
        ["docx", "markdown"],
        index=0 if config["output_format"] == "docx" else 1,
    )

    if st.button("Erweiterte Optionen"):
        st.session_state["advanced_open"] = not st.session_state.get("advanced_open", False)
    if st.session_state.get("advanced_open", False):
        with st.container():
            st.markdown("**Layout & Heuristiken**")
            column_detection = st.checkbox(
                "Spalten-Erkennung", value=config["column_detection"], help="Erkenne mehrere Spalten"
            )
            max_columns = st.number_input(
                "Max. Spalten", 1, 4, value=config["max_columns"], help="Maximale Spaltenanzahl"
            )
            h2 = st.number_input(
                "Heading2-Schwelle", 1.0, 3.0, value=float(config["heading_threshold_h2"]), step=0.1,
                help="Faktor × Medianhöhe"
            )
            h1 = st.number_input(
                "Heading1-Schwelle", 1.0, 3.0, value=float(config["heading_threshold_h1"]), step=0.1,
                help="Faktor × Medianhöhe"
            )
            centered = st.checkbox(
                "Zentriert", value=config["heading_extra_rules"]["centered"], help="Titel mittig"
            )
            all_caps = st.checkbox(
                "VERSALIEN", value=config["heading_extra_rules"]["all_caps"], help="Alle Großbuchstaben"
            )
            big_gap = st.checkbox(
                "Große Lücke", value=config["heading_extra_rules"]["big_gap"], help="Großer Abstand davor"
            )
            paragraph_gap = st.number_input(
                "Absatzlücke", 0.1, 3.0, value=float(config["paragraph_merge_gap"]), step=0.1,
                help="Max. Zeilenabstand in Medianhöhen"
            )
            indent_tol = st.number_input(
                "Einzugstoleranz", 0.0, 0.2, value=float(config["indent_tolerance"]), step=0.01,
                help="Toleranz für linke Kante"
            )
            hyphen_merge = st.checkbox(
                "Silbentrennung", value=config["hyphen_merge"], help="Trennstriche entfernen"
            )
            text_align = st.selectbox(
                "Textausrichtung", ["justify", "left", "center", "right"],
                index=["justify", "left", "center", "right"].index(config["text_alignment"])
            )
            keep_breaks = st.checkbox(
                "Zeilenumbrüche behalten", value=config["keep_line_breaks"], help="Originale Zeilenumbrüche"
            )
            debug_overlay = st.checkbox(
                "Debug-Overlay", value=config["debug_overlay"], help="BBox-Overlay speichern"
            )
            regex_cleanup = st.text_area(
                "Regex-Cleanup", value="\n".join(config["custom_regex_cleanup"]), help="Eine Regex je Zeile"
            )
    else:
        column_detection = config["column_detection"]
        max_columns = config["max_columns"]
        h2 = config["heading_threshold_h2"]
        h1 = config["heading_threshold_h1"]
        centered = config["heading_extra_rules"]["centered"]
        all_caps = config["heading_extra_rules"]["all_caps"]
        big_gap = config["heading_extra_rules"]["big_gap"]
        paragraph_gap = config["paragraph_merge_gap"]
        indent_tol = config["indent_tolerance"]
        hyphen_merge = config["hyphen_merge"]
        text_align = config["text_alignment"]
        keep_breaks = config["keep_line_breaks"]
        debug_overlay = config["debug_overlay"]
        regex_cleanup = "\n".join(config["custom_regex_cleanup"])

    updated = {
        "dpi": int(dpi),
        "languages": languages or config["languages"],
        "gpu": gpu,
        "column_detection": column_detection,
        "max_columns": int(max_columns),
        "heading_threshold_h2": float(h2),
        "heading_threshold_h1": float(h1),
        "heading_extra_rules": {"centered": centered, "all_caps": all_caps, "big_gap": big_gap},
        "paragraph_merge_gap": float(paragraph_gap),
        "indent_tolerance": float(indent_tol),
        "hyphen_merge": bool(hyphen_merge),
        "output_format": output_format,
        "text_alignment": text_align,
        "keep_line_breaks": bool(keep_breaks),
        "debug_overlay": bool(debug_overlay),
        "custom_regex_cleanup": [r for r in regex_cleanup.splitlines() if r.strip()],
    }
    return updated


# ---------------------------------------------------------------------------
# OCR pipeline
# ---------------------------------------------------------------------------

def ocr_pages(images: List[Image.Image], config: Dict[str, Any]) -> List[List[Dict[str, Any]]]:
    """Run EasyOCR on given images and return line dicts per page."""
    reader = easyocr.Reader(config["languages"], gpu=config["gpu"])
    pages: List[List[Dict[str, Any]]] = []
    for img in images:
        result = reader.readtext(np.array(img), detail=1, paragraph=False)
        lines = [{"bbox": r[0], "text": r[1], "conf": float(r[2])} for r in result]
        if config.get("column_detection") and len(lines) > 1:
            lines = _sort_columns(lines, img.size[0], config.get("max_columns", 2))
        pages.append(lines)
    return pages


def _sort_columns(lines: List[Dict[str, Any]], width: int, max_cols: int) -> List[Dict[str, Any]]:
    """Sort lines by columns using a simple KMeans clustering."""
    try:
        from sklearn.cluster import KMeans

        x_mid = [((min(p[0] for p in l["bbox"]) + max(p[0] for p in l["bbox"])) / 2) for l in lines]
        k = min(max_cols, len(lines))
        km = KMeans(n_clusters=k, n_init=10).fit(np.array(x_mid).reshape(-1, 1))
        labeled = list(zip(km.labels_, lines))
        labeled.sort(key=lambda t: (t[0], t[1]["bbox"][0][1]))
        return [l for _, l in labeled]
    except Exception:
        return sorted(lines, key=lambda l: l["bbox"][0][1])


def _pdf_to_images(pdf_file: bytes, dpi: int) -> List[Image.Image]:
    """Render PDF bytes to images. Tries pdf2image, falls back to pypdfium2."""
    images: List[Image.Image] = []
    try:
        from pdf2image import convert_from_bytes

        images = convert_from_bytes(pdf_file, dpi=dpi)
    except Exception:
        import pypdfium2 as pdfium

        pdf = pdfium.PdfDocument(io.BytesIO(pdf_file))
        for page in pdf:
            pil = page.render(scale=dpi / 72).to_pil()
            images.append(pil)
    return images


def _write_output(blocks_all: List[List[Dict[str, Any]]], config: Dict[str, Any]) -> str:
    """Write DOCX or Markdown output and return file path."""
    tmpdir = tempfile.mkdtemp()
    if config["output_format"] == "markdown":
        lines: List[str] = []
        for page in blocks_all:
            for block in page:
                if block["type"] == "h1":
                    lines.append("# " + block["text"])
                elif block["type"] == "h2":
                    lines.append("## " + block["text"])
                else:
                    lines.append(block["text"])
                lines.append("")
        md = "\n".join(lines)
        path = os.path.join(tmpdir, "output.md")
        with open(path, "w", encoding="utf-8") as f:
            f.write(md)
        return path
    else:
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except Exception:
            # Fallback to markdown
            return _write_output(blocks_all, {**config, "output_format": "markdown"})
        doc = Document()
        for page in blocks_all:
            for block in page:
                if block["type"] == "h1":
                    doc.add_heading(block["text"], level=1)
                elif block["type"] == "h2":
                    doc.add_heading(block["text"], level=2)
                else:
                    p = doc.add_paragraph(block["text"])
                    if config["text_alignment"] == "justify":
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        path = os.path.join(tmpdir, "output.docx")
        doc.save(path)
        return path


def run_ocr(pdf_file: bytes, config: Dict[str, Any]) -> Dict[str, Any]:
    """Execute OCR pipeline and return results and metrics."""
    start = time.time()
    images = _pdf_to_images(pdf_file, config["dpi"])
    if not images:
        return {"pages": 0, "lines": 0, "avg_conf": 0.0}
    page_lines = ocr_pages(images, config)
    blocks_all: List[List[Dict[str, Any]]] = []
    page_stats: List[Dict[str, Any]] = []
    total_lines = 0
    confidences: List[float] = []
    classified_pages: List[List[Dict[str, Any]]] = []
    for img, lines in zip(images, page_lines):
        classified, median = heur.classify_lines(lines, img.size, config)
        blocks = heur.build_blocks(classified, median, img.size, config)
        blocks_all.append(blocks)
        page_stats.append(
            {
                "median_height": median,
                "line_heights": [l["height"] for l in classified],
                "h1": sum(1 for l in classified if l["level"] == "h1"),
                "h2": sum(1 for l in classified if l["level"] == "h2"),
            }
        )
        total_lines += len(lines)
        confidences.extend([l["conf"] for l in lines])
        classified_pages.append(classified)
    if config.get("custom_regex_cleanup"):
        import re

        patterns = [re.compile(r) for r in config["custom_regex_cleanup"]]
        for page in blocks_all:
            for block in page:
                text = block["text"]
                for pat in patterns:
                    text = pat.sub("", text)
                block["text"] = text
    out_path = _write_output(blocks_all, config)
    runtime = time.time() - start
    logging.info("Processed %s pages in %.2fs", len(images), runtime)
    return {
        "pages": len(images),
        "lines": total_lines,
        "avg_conf": float(np.mean(confidences)) if confidences else 0.0,
        "blocks": blocks_all,
        "page_stats": page_stats,
        "images": images,
        "classified_lines": classified_pages,
        "doc_path": out_path,
    }


# ---------------------------------------------------------------------------
# Reporting
# ---------------------------------------------------------------------------

def make_debug_plots(results: Dict[str, Any], config: Dict[str, Any]) -> Dict[str, str]:
    """Create diagnostic plots and return their file paths."""
    paths: Dict[str, str] = {}
    tmpdir = tempfile.mkdtemp()
    paths["dir"] = tmpdir
    try:
        import matplotlib.pyplot as plt
    except Exception:
        # Create placeholder images if matplotlib is unavailable
        img = Image.new("RGB", (200, 100), "white")
        img.save(os.path.join(tmpdir, "placeholder.png"))
        paths["histogram"] = os.path.join(tmpdir, "placeholder.png")
        paths["heading_counts"] = os.path.join(tmpdir, "placeholder.png")
        return paths
    if results.get("page_stats"):
        stats0 = results["page_stats"][0]
        heights = stats0["line_heights"]
        if heights:
            plt.figure()
            plt.hist(heights, bins=20, color="gray")
            med = stats0["median_height"]
            plt.axvline(med, color="blue", label="median")
            plt.axvline(med * config["heading_threshold_h2"], color="orange", label="H2")
            plt.axvline(med * config["heading_threshold_h1"], color="red", label="H1")
            plt.legend()
            path = os.path.join(tmpdir, "line_heights.png")
            plt.savefig(path)
            plt.close()
            paths["histogram"] = path
    if results.get("page_stats"):
        h1 = [s["h1"] for s in results["page_stats"]]
        h2 = [s["h2"] for s in results["page_stats"]]
        x = np.arange(len(h1))
        plt.figure()
        plt.bar(x - 0.2, h1, width=0.4, label="H1")
        plt.bar(x + 0.2, h2, width=0.4, label="H2")
        plt.xlabel("Seite")
        plt.ylabel("Anzahl")
        plt.legend()
        path2 = os.path.join(tmpdir, "heading_counts.png")
        plt.savefig(path2)
        plt.close()
        paths["heading_counts"] = path2
    if config.get("debug_overlay") and results.get("images"):
        img = results["images"][0].copy()
        draw = ImageDraw.Draw(img)
        for line in results["classified_lines"][0]:
            color = "green"
            if line["level"] == "h1":
                color = "blue"
            elif line["level"] == "h2":
                color = "red"
            draw.polygon([tuple(p) for p in line["bbox"]], outline=color)
        overlay_path = os.path.join(tmpdir, "overlay.png")
        img.save(overlay_path)
        paths["overlay"] = overlay_path
    return paths


def generate_report(results: Dict[str, Any], config: Dict[str, Any], plots: Dict[str, str]) -> str:
    """Generate an HTML report and return its file path."""
    tmpdir = plots.get("dir", tempfile.mkdtemp())
    rows = "".join(
        f"<tr><th>{k}</th><td>{v}</td></tr>" for k, v in config.items()
    )
    images_html = "".join(
        f'<h3>{name}</h3><img src="{os.path.basename(path)}" alt="{name}">' for name, path in plots.items() if name != "dir"
    )
    html = f"""
    <html><body>
    <h1>Tuning Report</h1>
    <table>{rows}</table>
    {images_html}
    </body></html>
    """
    path = os.path.join(tmpdir, "report.html")
    with open(path, "w", encoding="utf-8") as f:
        f.write(html)
    return path


# ---------------------------------------------------------------------------
# Streamlit UI
# ---------------------------------------------------------------------------

def main() -> None:
    st.set_page_config(page_title="PDF OCR")
    st.title("PDF OCR Pipeline")
    st.write("Verarbeite PDFs mit anpassbarer OCR und erhalte einen Tuning-Report.")

    uploaded = st.file_uploader("PDF hochladen", type="pdf")
    if uploaded:
        st.session_state["uploaded_pdf"] = uploaded.getvalue()

    config = st.session_state.get("config", build_default_config())
    config = render_controls(config)
    st.session_state["config"] = config

    if st.button("OCR ausführen"):
        if not st.session_state.get("uploaded_pdf"):
            st.warning("Bitte zuerst eine PDF-Datei hochladen.")
        else:
            with st.spinner("Verarbeite..."):
                results = run_ocr(st.session_state["uploaded_pdf"], config)
                plots = make_debug_plots(results, config)
                report_path = generate_report(results, config, plots)
                st.session_state["results"] = results
                st.session_state["plots"] = plots
                st.session_state["report_path"] = report_path

    if st.session_state.get("results"):
        res = st.session_state["results"]
        st.success(
            f"Seiten: {res['pages']} | Zeilen: {res['lines']} | Ø Konfidenz: {res['avg_conf']:.2f}"
        )
        for name, path in st.session_state["plots"].items():
            if name == "dir":
                continue
            st.image(path, caption=name)
        with open(res["doc_path"], "rb") as f:
            st.download_button(
                "Ergebnis herunterladen",
                f,
                file_name=os.path.basename(res["doc_path"]),
            )
        with open(st.session_state["report_path"], "rb") as f:
            st.download_button(
                "Report herunterladen",
                f,
                file_name="report.html",
            )


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    main()
